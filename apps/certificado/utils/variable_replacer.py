"""
Reemplazo de variables en documentos Word.

Usa python-docx para reemplazar placeholders {{VARIABLE}} con valores reales.
"""

import re
import logging
from typing import Dict
from docx import Document
from docx.shared import Pt


logger = logging.getLogger(__name__)


class VariableReplacer:
    """
    Clase para reemplazar variables en documentos Word (.docx).
    
    Variables soportadas (formato {{VARIABLE}}):
        {{NOMBRES}}, {{MODALIDAD}}, {{NOMBRE_EVENTO}}, {{DURACION}},
        {{FECHA_INICIO}}, {{FECHA_FIN}}, {{TIPO}}, {{TIPO_EVENTO}},
        {{FECHA_EMISION}}, {{OBJETIVO_PROGRAMA}}, {{CONTENIDO}}
    """
    
    # Regex pattern para detectar variables (incluyendo espacios)
    VARIABLE_PATTERN = re.compile(r'\{\{([A-Z_ ]+)\}\}')
    
    @staticmethod
    def replace_in_document(doc_path: str, variables: Dict[str, str]) -> Document:
        """
        Carga un documento Word y reemplaza todas las variables.
        
        Args:
            doc_path: Ruta al archivo .docx de plantilla
            variables: Diccionario de reemplazos
                Ejemplo: {"NOMBRES": "Juan Pérez", "MODALIDAD": "Virtual", ...}
        
        Returns:
            Objeto Document modificado (listo para guardar)
        
        Ejemplo:
            >>> from apps.certificado.utils.variable_replacer import VariableReplacer
            >>> variables = {
            ...     "NOMBRES": "Juan Pérez",
            ...     "MODALIDAD": "Virtual",
            ...     "NOMBRE_EVENTO": "Taller de Python"
            ... }
            >>> doc = VariableReplacer.replace_in_document('/path/to/template.docx', variables)
            >>> doc.save('/path/to/output.docx')
        """
        try:
            # Cargar documento
            doc = Document(doc_path)
            logger.info(f"Documento cargado: {doc_path}")
            
            # Normalizar variables (asegurar mayúsculas)
            variables_upper = {k.upper(): v for k, v in variables.items()}
            
            # Reemplazar en párrafos
            VariableReplacer._replace_in_paragraphs(doc.paragraphs, variables_upper)
            
            # Reemplazar en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        VariableReplacer._replace_in_paragraphs(cell.paragraphs, variables_upper)
            
            # Reemplazar en headers y footers
            for section in doc.sections:
                VariableReplacer._replace_in_paragraphs(section.header.paragraphs, variables_upper)
                VariableReplacer._replace_in_paragraphs(section.footer.paragraphs, variables_upper)
            
            logger.info(f"Variables reemplazadas exitosamente")
            return doc
            
        except Exception as e:
            logger.error(f"Error al reemplazar variables en documento: {str(e)}")
            raise
    
    @staticmethod
    def _replace_in_paragraphs(paragraphs, variables: Dict[str, str]):
        """
        Reemplaza variables en una lista de párrafos.
        
        Args:
            paragraphs: Lista de párrafos de python-docx
            variables: Diccionario de reemplazos
        """
        for paragraph in paragraphs:
            VariableReplacer._replace_in_paragraph(paragraph, variables)
    
    @staticmethod
    def _replace_in_paragraph(paragraph, variables: Dict[str, str]):
        """
        Reemplaza variables en un solo párrafo, preservando el formato.
        
        Args:
            paragraph: Párrafo de python-docx
            variables: Diccionario de reemplazos
        """
        # Obtener todo el texto del párrafo
        full_text = paragraph.text
        
        # Si no hay variables, salir
        if '{{' not in full_text:
            return
        
        # Buscar todas las variables en el texto
        variables_found = VariableReplacer.VARIABLE_PATTERN.findall(full_text)
        
        if not variables_found:
            return
        
        # Reemplazar cada variable
        new_text = full_text
        for var_name in variables_found:
            placeholder = f'{{{{{var_name}}}}}'
            
            if var_name in variables:
                # Reemplazar con valor
                replacement = str(variables[var_name])
                new_text = new_text.replace(placeholder, replacement)
                logger.debug(f"Reemplazado: {placeholder} → {replacement}")
            else:
                # Variable no encontrada, dejar como está y loggear
                logger.warning(f"Variable no encontrada: {placeholder}")
        
        # Si hubo cambios, actualizar el párrafo
        if new_text != full_text:
            VariableReplacer._update_paragraph_text(paragraph, new_text)
    
    @staticmethod
    def _update_paragraph_text(paragraph, new_text: str):
        """
        Actualiza el texto de un párrafo preservando el formato del primer run.
        
        Args:
            paragraph: Párrafo de python-docx
            new_text: Nuevo texto
        """
        # Guardar formato del primer run si existe
        first_run_font = None
        if paragraph.runs:
            first_run_font = paragraph.runs[0].font
        
        # Limpiar todos los runs
        for run in paragraph.runs:
            run.text = ''
        
        # Eliminar runs vacíos
        paragraph.clear()
        
        # Agregar nuevo run con el texto actualizado
        new_run = paragraph.add_run(new_text)
        
        # Aplicar formato del run original si existía
        if first_run_font:
            new_run.font.name = first_run_font.name
            new_run.font.size = first_run_font.size
            new_run.font.bold = first_run_font.bold
            new_run.font.italic = first_run_font.italic
            new_run.font.underline = first_run_font.underline
            new_run.font.color.rgb = first_run_font.color.rgb


def replace_variables_in_template(template_path: str, variables: Dict[str, str]) -> Document:
    """
    Función helper para reemplazar variables en una plantilla.
    
    Args:
        template_path: Ruta al archivo .docx
        variables: Diccionario de variables a reemplazar
    
    Returns:
        Documento modificado
    
    Ejemplo:
        >>> doc = replace_variables_in_template('/path/to/template.docx', {
        ...     "NOMBRES": "Juan Pérez",
        ...     "DURACION": "40 horas"
        ... })
        >>> doc.save('/path/to/output.docx')
    """
    return VariableReplacer.replace_in_document(template_path, variables)
